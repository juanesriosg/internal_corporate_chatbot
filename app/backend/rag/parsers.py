from __future__ import annotations

import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from app.backend.rag.models import DocumentRecord


def parse_document(path: Path, document: DocumentRecord) -> str:
    parser = document.format.lower()
    if parser == "markdown":
        return _strip_markdown_frontmatter(path.read_text(encoding="utf-8"))
    if parser == "txt":
        return path.read_text(encoding="utf-8")
    if parser == "html":
        return _parse_html(path)
    if parser == "pdf":
        return _parse_pdf(path)
    if parser == "docx":
        return _parse_docx(path)
    raise ValueError(f"Unsupported document format '{document.format}' for {path}")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_markdown_frontmatter(text: str) -> str:
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            text = parts[2]
    return normalize_text(text)


def _parse_html(path: Path) -> str:
    soup = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    body = soup.get_text("\n", strip=True)
    return normalize_text(f"{title}\n\n{body}" if title else body)


def _parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"Page {page_number}\n{text}")
    return normalize_text("\n\n".join(pages))


def _parse_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return normalize_text("\n\n".join(parts))
